from tkinter import *
import tkinter.scrolledtext as st
import spacy_final
from tkinter.filedialog import asksaveasfilename
import tkinter.messagebox
import docx


# error message if user has not entered a search word
def no_word():
    global no_word_message
    no_word_message = Toplevel()
    no_word_message.title("No search term specified")
    no_word_message.geometry("500x100")
    Label(no_word_message, text="No Search term specified! Please enter a word and try again.",
          fg="red", font=('arial', 12, "bold")).pack()
    Button(no_word_message, text="Close", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=close_no_word).pack()


# checks if the user has entere a search word or not
def topic_check(word):
    if not word:
        no_word()
    else:
        view_topic(word)


def view_topic(word):
    global topic_viewer
    topic_viewer = Toplevel()
    topic_viewer.title("Search Topic")
    topic_viewer.geometry("1000x970")
    topic_viewer.config(bg="white")
    global text
    text = spacy_final.main(word)
    Label(topic_viewer, text=topic.get(), font=('arial', 20, 'bold'), relief="groove", fg="white",
          bg="green", width=300).pack()
    text_area = st.ScrolledText(topic_viewer,
                                width=400,
                                height=40,
                                font=("Calibri",
                                      11))

    text_area.pack()
    text_area.insert(INSERT, text[1])
    text_area.configure(state='disabled', wrap=WORD)

    Button(topic_viewer, text="Save as Word", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=lambda: save_word(text[1])).pack(pady=5)

    Button(topic_viewer, text="Save as TXT", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=lambda: save_txt(text[1])).pack()

    Button(topic_viewer, text="View Sources", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=view_sources).pack(pady=5)

    Button(topic_viewer, text="View Statistics", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=view_stats).pack()

    Button(topic_viewer, text="Close", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=close_topic).pack(pady=5)


# displays statistics from the text (common words and named enitities)
def view_stats():
    global stats_viewer
    stats_viewer = Toplevel()
    stats_viewer.title("Search Topic")
    stats_viewer.geometry("700x800")
    stats_viewer.config(bg="white")

    word_freq = '\n'.join(str(txt) for txt in text[2])
    ner = '\n'.join(str(txt) for txt in text[3])

    Label(stats_viewer, text="Word frequencies", font=('arial', 20, 'bold'), relief="raised", fg="white",
          bg="green").pack()

    displ_stats = st.ScrolledText(stats_viewer, font=('arial', 10), height=20)
    displ_stats.insert(INSERT, word_freq)
    displ_stats.config(state=DISABLED)
    displ_stats.pack()

    Label(stats_viewer, text="Named Enitities", font=('arial', 20, 'bold'), relief="raised", fg="white",
          bg="green").pack()

    displ_stats2 = st.ScrolledText(stats_viewer, font=('arial', 10), height=20)
    displ_stats2.insert(INSERT, ner)
    displ_stats2.config(state=DISABLED)
    displ_stats2.pack()

    Button(stats_viewer, text="Close", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=close_stats).pack()


# displaying the sources the text was extracted from
def view_sources():
    global sources_viewer
    sources_viewer = Toplevel()
    sources_viewer.title("Search Topic")
    sources_viewer.geometry("500x500")
    sources_viewer.config(bg="white")

    sources = '\n\n'.join(text[0])

    displ_sources = Text(sources_viewer, font=('arial', 10))
    displ_sources.insert(INSERT, sources)
    displ_sources.config(state=DISABLED)
    displ_sources.pack()

    Button(sources_viewer, text="Close", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=close_sources).pack()


# saving extracted text as a .txt file
def save_txt(text_to_save):
    path_to_pref = asksaveasfilename(
        defaultextension='.txt', filetypes=[("Text files", '*.txt')],
        title="Choose filename")
    with open(path_to_pref, 'wb') as f:
        f.write(text_to_save.encode('utf-8'))
        f.close()


# saving extracted text as a .docx file
def save_word(text_to_save):
    path_to_pref = asksaveasfilename(
        defaultextension='.docx', filetypes=[("Word files", '*.docx')],
        title="Choose filename")
    mydoc = docx.Document()
    mydoc.add_heading("{}".format(topic.get()), 0)
    mydoc.add_paragraph(text_to_save)
    mydoc.save(path_to_pref)


# exit function
def Exit():
    wayOut = tkinter.messagebox.askyesno(
        "Find It All", "Do you want to exit?")
    if wayOut > 0:
        root.destroy()
        return


def main_menu():
    global root
    root = Tk()
    root.config(bg="white")
    root.title("Main Menu")
    root.geometry("500x500")
    Label(root, text='Welcome to Find It All',  bd=20, font=('arial', 20, 'bold'), relief="raised", fg="white",
          bg="green", width=300).pack()
    Label(root, text="").pack()

    Button(root, text='Search Topic', height="1", width="20", bd=8, font=('arial', 12, 'bold'), relief="raised", fg="white",
           bg="green", command=search_topic).pack()
    Label(root, text="").pack()

    Button(root, text='Exit', height="1", width="20", bd=8, font=('arial', 12, 'bold'), relief="raised", fg="white",
           bg="green", command=Exit).pack()
    Label(root, text="").pack()


def search_topic():
    global topic_searcher
    topic_searcher = Toplevel()
    topic_searcher.title("Search Topic")
    topic_searcher.geometry("480x300")
    topic_searcher.config(bg="white")

    Label(topic_searcher, text='Please Enter a keyword you would like to get information for', bd=5, font=('arial', 12, 'bold'), relief="raised", fg="white",
          bg="green", width=300).pack()

    global topic
    topic = StringVar()
    Entry(topic_searcher, textvariable=topic).pack()

    Button(topic_searcher, text="Search", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=lambda: topic_check(topic.get())).pack(pady=5)

    Label(root, text="").pack()

    Button(topic_searcher, text="Close", bg="green", fg='white', relief="raised",
           font=('arial', 12, 'bold'), command=close_search).pack(pady=5)


# close window functions
def close_search():
    topic_searcher.destroy()


def close_sources():
    sources_viewer.destroy()


def close_stats():
    stats_viewer.destroy()


def close_topic():
    topic_viewer.destroy()


def close_no_word():
    no_word_message.destroy()


main_menu()
root.mainloop()
